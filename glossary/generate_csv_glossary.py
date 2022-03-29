import csv
import re
from pathlib import Path

from botok import Text
from docx import Document
from pyewts import pyewts

from .format_unicode import bold, ital

converter = pyewts()


def generate_csv(in_path, out_file):
    parsed = parse_docx(in_path)
    rows = parsed_2_rows(parsed)
    with open(out_file, 'w', newline='') as csvfile:
        writer = csv.writer(csvfile, delimiter=',', quotechar='|')
        for row in rows:
            writer.writerow(row)


def parsed_2_rows(parsed):
    rows = [['term', 'pos', 'comment', 'is_case_sensitive', 'fr']]
    for lemma, fields in parsed.items():
        cur_row = [''] * 5
        num, word = lemma
        # term column: lemma
        cur_row[0] = word

        # pos column: ignored

        # comment column: words, entry number, definition and notes
        defn = []
        words = ' '.join([f for f in fields['words'] if f]).replace(',', ' -')
        defn.append(ital(words))
        defn.append(f'{num}. ')
        for section in ['def', 'notes']:
            if ''.join(fields[section]):
                sec = ' '.join([f for f in fields[section] if f]).replace(',', '-')
                defn.append(f'{bold(section)}: {sec}')
        defn = ' '.join(defn)
        cur_row[2] = defn

        # fr column: url
        if len(fields['url']) < 500:
            cur_row[4] = fields['url']

        rows.append(cur_row)
    return rows


def parse_docx(in_path):
    def add_section(num, section):
        num += 1
        par = doc.paragraphs[num]

        if section not in gloss[cur_entry]:
            gloss[cur_entry][section] = ['']

        while par.style.name != 'Heading 4':
            gloss[cur_entry][section].append(par.text)
            num += 1
            par = doc.paragraphs[num]
        num -= 1
        return num

    files = sorted(list(Path(in_path).glob('*.docx')))
    gloss = {}
    for file in files:
        print(file.name)
        doc = Document(file)
        cur_entry = None
        num = 0
        while num < len(doc.paragraphs):
            par = doc.paragraphs[num]
            if not par.text:
                num += 1
                continue

            if par.style.name == 'Heading 2':
                e_num, word = par.text.split(' ', 1)
                word = segment_in_words(word)
                cur_entry = tuple([e_num, word])
                if cur_entry[0] and cur_entry not in gloss:
                    gloss[cur_entry] = {}
                    # add christian steinert url
                    url = 'https://dictionary.christian-steinert.de/#%7B%22activeTerm%22%3A%22{word}%22%2C%22lang%22%3A%22tib%22%2C%22inputLang%22%3A%22tib%22%2C%22currentListTerm%22%3A%22{word}%22%2C%22forceLeftSideVisible%22%3Afalse%2C%22offset%22%3A0%7D'
                    wylie = converter.toWylie(cur_entry[1])
                    wylie = wylie.replace('_', '').replace(' ', '%20')
                    url = url.format(word=wylie)
                    gloss[cur_entry]['url'] = url
            elif par.style.name == 'Heading 4':
                if par.text == 'Termes utilisés':
                    num = add_section(num, 'words')
                elif par.text == 'Définition':
                    num = add_section(num, 'def')
                elif par.text == 'Notes':
                    num = add_section(num, 'notes')
            num += 1
    return gloss


def segment_in_words(string):
    string = string.strip().rstrip('།').rstrip('་')
    t = Text(string)
    tokenized = t.tokenize_words_raw_text
    # format tokens
    tokenized = re.sub('([^།་_]) ([^_།་])', '\g<1>␣\g<2>', tokenized)  # affixed particles
    tokenized = re.sub('_', ' ', tokenized)  # spaces
    return tokenized
