from pathlib import Path
from textwrap import dedent
from csv import DictReader

import yaml
import docx
from docx import Document
import pyewts
from botok import TokChunks
from tibetan_sort import TibetanSort

converter = pyewts.pyewts()
sorter = TibetanSort()


def parse_glossaries(in_path, out_path):
    # parse txt glossaries
    joined = {}
    gloss_txt = Path(in_path) / 'raw_glossaries'
    for f in gloss_txt.glob('*.txt'):
        print(f.name)
        parse_bar_separated(f, joined)

    # todo: parse spreadsheet glossary
    gloss_csv = Path(in_path) / 'spreadsheets'
    for f in gloss_csv.glob('*.csv'):
        print(f.name)
        parse_csv(f, joined)

    # sort content
    total = {}
    sorted_words = sorter.sort_list(joined.keys())
    for num, word in enumerate(sorted_words):
        entry = joined[word]
        sorted_entry = [(k, entry[k]) for k in sorted(entry.keys())]
        total[num+1] = (word, sorted_entry)

    # export in docx
    export_docx(total, out_path)


def parse_csv(in_file, joined):
    with in_file.open(newline='') as csvfile:
        reader = DictReader(csvfile)
        for row in reader:
            word = row['Tibétain'].strip()
            if not word or word[0].isupper():
                continue
            others = ''
            if '\n' in word:
                word, others = word.split('\n', maxsplit=1)
            others = others.strip()
            word = converter.toUnicode(word)
            word = '་'.join(TokChunks(word).get_syls())
            word = add_shad(word)
            skrt = row['Sanskrit'].strip()
            root_en = row['Anglais'].strip()
            root_fr = row['Sens racine Français'].strip()
            alter = row['Autres termes rencontrés'].strip()
            entry = []
            if others:
                entry.append(others)
            if skrt:
                entry.append(skrt)
            roots = [r for r in [root_fr, root_en] if r]
            if roots:
                entry.append(f'Sens racine: {", ".join(roots)}')
            if alter:
                entry.append(f'Alternatives: {alter}')
            entry = '\n'.join(entry)

            if word not in joined:
                joined[word] = {}
            joined[word][in_file.stem] = [entry]


def parse_bar_separated(in_file, joined):
    sep = '|'
    dump = in_file.read_text()
    name = in_file.stem

    # sanity check
    mal_formed = []

    for num, line in enumerate(dump.split('\n')):
        if not line.strip():
            continue

        if line.count(sep) != 1:
            mal_formed.append((num+1, line))
        else:
            word, defnt = [e.strip() for e in line.split(sep)]

            # cleanup (TokChunks) and convert to unicode
            word = converter.toUnicode(word)
            word = '་'.join(TokChunks(word).get_syls())
            word = add_shad(word)

            if word not in joined:
                joined[word] = {}

            if name not in joined[word]:
                joined[word][name] = []

            joined[word][name].append(defnt)
    if mal_formed:
        print('\n'.join([': '.join(m) for m in mal_formed]))
        exit(1)


def add_shad(word):
    if word.endswith('ང'):
        return word + '་།'
    elif word[-1] in ['ཀ', 'ག', 'ཤ']:
        return word
    else:
        return word + '།'


def export_docx(data, out_path):
    def add_entries(current_entry):
        word, entries = data[current_entry][0], data[current_entry][1]
        doc.add_heading(f'{current_entry} {word}', 3)
        for name, entry in entries:
            doc.add_heading(name, 4)
            for n, defnt in enumerate(entry):
                doc.add_paragraph(f'- {defnt}')

        url = 'https://dictionary.christian-steinert.de/#%7B%22activeTerm%22%3A%22{word}%22%2C%22lang%22%3A%22tib%22%2C%22inputLang%22%3A%22tib%22%2C%22currentListTerm%22%3A%22{word}%22%2C%22forceLeftSideVisible%22%3Afalse%2C%22offset%22%3A0%7D'
        wylie = converter.toWylie(word)
        p = doc.add_paragraph()
        add_hyperlink(p, url.format(word=wylie), 'Christian Steinert', '#0000EE', True)

    config = parse_config(data)
    current_entry = 1
    vol_count = 1
    for start, end in config['start_ends']:
        doc = Document()
        while current_entry <= len(data) and data[current_entry][0] != end:
            add_entries(current_entry)
            current_entry += 1
        else:
            if current_entry <= len(data):
                add_entries(current_entry)
                current_entry += 1

        s_name = start.split('་')
        s_name = '་'.join(s_name) if len(s_name) <= 5 else '་'.join(s_name[:5]) + '[...]'
        e_name = end.split('་')
        e_name = '་'.join(e_name) if len(e_name) <= 5 else '་'.join(e_name[:5]) + '[...]'
        out_file = Path(out_path) / f'{vol_count} {s_name} — {e_name}.docx'
        doc.save(out_file)
        vol_count += 1


def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


def parse_config(data):
    conf_file = Path('config.yaml')
    if not conf_file.is_file():
        template = dedent("""\
        entries_per_file: 1000
        start_ends:
        """)
        conf_file.write_text(yaml.safe_dump(template, allow_unicode=True))

    conf = yaml.safe_load(conf_file.read_text())
    if not conf['start_ends']:
        starting_idx = list(range(1, len(data), conf['entries_per_file']))
        starting_entries = [data[s][0] for s in starting_idx]
        ending_idx = list(range(conf['entries_per_file'], len(data), conf['entries_per_file'])) + [len(data)]
        ending_entries = [data[e][0] for e in ending_idx]
        files = [[s, e] for s, e in zip(starting_entries, ending_entries)]
        conf['start_ends'] = files
        conf_file.write_text(yaml.safe_dump(conf, allow_unicode=True))
    return conf


parse_glossaries('../content', '../content/')
