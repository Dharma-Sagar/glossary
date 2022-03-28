import json
from pathlib import Path

import yaml
import docx
from docx import Document
import pyewts
from tibetan_sort import TibetanSort

converter = pyewts.pyewts()
sorter = TibetanSort()


def export_docx(json_file, out_path):
    def add_entries(current_entry):
        word, entries = data[str(current_entry)][0], data[str(current_entry)][1]
        url = 'https://dictionary.christian-steinert.de/#%7B%22activeTerm%22%3A%22{word}%22%2C%22lang%22%3A%22tib%22%2C%22inputLang%22%3A%22tib%22%2C%22currentListTerm%22%3A%22{word}%22%2C%22forceLeftSideVisible%22%3Afalse%2C%22offset%22%3A0%7D'
        wylie = converter.toWylie(word)

        doc.add_heading(f'{current_entry} {word}', 2)
        doc.add_heading('Termes utilisés', 4)
        doc.add_heading('Définition', 4)
        doc.add_heading('Notes', 4)
        doc.add_heading('À consulter', 4)
        par = doc.add_paragraph()
        add_hyperlink(par, url.format(word=wylie), 'Christian Steinert', '#0000EE', True)
        par.add_run().add_break()

        e_num = 1
        for name, entry in entries:
            if e_num > 1:
                name = f'\n{name}'
            run = par.add_run(name + ' ')
            run.font.bold = True
            run.font.italic = True

            d_num = 1
            for n, defnt in enumerate(entry):
                if d_num > 1:
                    defnt = f'\n⁃ {defnt}'
                else:
                    defnt = f'⁃ {defnt}'
                par.add_run(defnt)

                d_num += 1
            e_num += 1

    data = json.loads(Path(json_file).read_text())
    config = parse_config(data)
    current_entry = 1
    vol_count = 1
    for start, end in config['start_ends']:
        doc = Document()
        while current_entry <= len(data) and data[str(current_entry)][0] != end:
            add_entries(current_entry)
            current_entry += 1
        else:
            if current_entry <= len(data):
                add_entries(current_entry)
                current_entry += 1

        s_name = start.split('་')
        s_name = '་'.join(s_name) if len(s_name) <= 5 else '་'.join(s_name[:5]) + '[...]'
        e_name = end.split('་')
        e_name = '་'.join(e_name) if len(e_name) <= 5 else '་'.join(e_name[:5]) + '[...]'
        out_file = Path(out_path) / f'{vol_count} {s_name} — {e_name}.docx'
        doc.save(out_file)
        vol_count += 1


def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


def parse_config(data):
    conf_file = Path('config.yaml')
    if not conf_file.is_file():
        template = {'entries_per_file': 700, 'start_ends': ''}
        conf_file.write_text(yaml.safe_dump(template))

    conf = yaml.safe_load(conf_file.read_text())
    if not conf['start_ends']:
        starting_idx = list(range(1, len(data), conf['entries_per_file']))
        starting_entries = [data[str(s)][0] for s in starting_idx]
        ending_idx = list(range(conf['entries_per_file'], len(data), conf['entries_per_file'])) + [len(data)]
        ending_entries = [data[str(e)][0] for e in ending_idx]
        files = [[s, e] for s, e in zip(starting_entries, ending_entries)]
        conf['start_ends'] = files
        conf_file.write_text(yaml.safe_dump(conf, allow_unicode=True))
    return conf
